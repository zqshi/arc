"""文档上传、文本提取和Feature解析服务。"""

from __future__ import annotations

import logging
import uuid
from io import BytesIO

from sqlalchemy.ext.asyncio import AsyncSession

from arc.domain.errors import AppError
from arc.domain.planning.entity import Document
from arc.infrastructure.repositories.planning import DocumentRepository
from arc.infrastructure.storage import get_storage

logger = logging.getLogger(__name__)

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "text/markdown",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


def _sanitize_filename(filename: str) -> str:
    from pathlib import PurePosixPath
    name = PurePosixPath(filename).name
    if not name or name in (".", ".."):
        raise AppError("Invalid filename")
    return name


def _storage_key(project_id: uuid.UUID, doc_id: uuid.UUID, filename: str) -> str:
    safe_name = _sanitize_filename(filename)
    return f"documents/{project_id}/{doc_id}/{safe_name}"


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.doc_repo = DocumentRepository(db)

    async def upload(
        self,
        project_id: uuid.UUID,
        filename: str,
        content_type: str,
        data: bytes,
    ) -> Document:
        if content_type not in ALLOWED_CONTENT_TYPES:
            raise AppError(f"不支持的文件类型: {content_type}")
        if len(data) > MAX_FILE_SIZE:
            raise AppError("文件大小超过20MB限制")

        doc = Document(
            project_id=project_id,
            filename=filename,
            content_type=content_type,
            size=len(data),
        )

        storage = get_storage()
        key = _storage_key(project_id, doc.id, filename)
        await storage.async_upload(key, data, content_type, max_size=MAX_FILE_SIZE)
        doc.storage_path = key

        doc.mark_processing()
        await self.doc_repo.create(doc)

        try:
            text = await self._extract_text(data, content_type)
            features = await self._parse_features(text, filename, project_id=project_id)
            doc.mark_ready(text, features)
        except Exception as exc:
            logger.exception("Document processing failed: %s", exc)
            doc.mark_error()

        await self.doc_repo.update(doc)
        return doc

    async def list_by_project(
        self, project_id: uuid.UUID, *, skip: int = 0, limit: int = 100
    ) -> list[Document]:
        return await self.doc_repo.list_by_project(project_id, skip=skip, limit=limit)

    async def delete(self, doc_id: uuid.UUID) -> bool:
        doc = await self.doc_repo.get_by_id(doc_id)
        if doc and doc.storage_path:
            storage = get_storage()
            await storage.async_delete(doc.storage_path)
        return await self.doc_repo.delete(doc_id)

    async def _extract_text(self, data: bytes, content_type: str) -> str:
        if content_type in ("text/markdown", "text/plain"):
            return data.decode("utf-8", errors="replace")
        if content_type == "application/pdf":
            return await self._extract_pdf(data)
        if "wordprocessingml" in content_type:
            return await self._extract_docx(data)
        return ""

    @staticmethod
    async def _extract_pdf(data: bytes) -> str:
        try:
            import pypdf

            reader = pypdf.PdfReader(BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("pypdf not installed, PDF text extraction unavailable")
            return "[PDF文件 - 需安装pypdf库进行文本提取]"

    @staticmethod
    async def _extract_docx(data: bytes) -> str:
        try:
            import docx

            doc = docx.Document(BytesIO(data))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed, DOCX extraction unavailable")
            return "[DOCX文件 - 需安装python-docx库进行文本提取]"

    async def _parse_features(
        self, text: str, filename: str, *, project_id: uuid.UUID | None = None
    ) -> list[dict]:
        """用AI从文档文本中提取功能点列表。"""
        if not text or len(text.strip()) < 50:
            return []

        from arc.application.ai.adapter_pool import adapter_pool
        from arc.application.ai.json_extract import extract_json
        from arc.application.ai.llm_adapter import LLMMessage
        from arc.application.llm.service import LLMProviderService

        prompt = f"""从以下文档中提取所有可独立交付的功能需求。

## 文档来源: {filename}
## 文档内容:
{text[:8000]}

输出 JSON 数组:
```json
[
  {{
    "title": "功能名称",
    "description": "功能描述",
    "complexity": "S/M/L/XL",
    "category": "所属模块",
    "priority_hint": "high/medium/low",
    "dependencies": ["依赖的其他功能title"]
  }}
]
```"""

        llm_config = await LLMProviderService.resolve_for_context(
            self.db, project_id=project_id
        )
        async with adapter_pool.acquire_for_project(llm_config) as adapter:
            response = await adapter.chat([LLMMessage(role="user", content=prompt)])

        result = extract_json(response.content)
        if isinstance(result, list):
            return result
        return []
